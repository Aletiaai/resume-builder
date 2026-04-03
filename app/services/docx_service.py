"""DOCX generation service.

Implements the exact formatting spec from .claude/skills/resume-tailor/references/formatting.md.
Uses python-docx. Returns document bytes — never writes to local disk.
"""

import io
import logging
import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

from app.models.schemas import TailoredResume

logger = logging.getLogger(__name__)

# Color constants
COLOR_TEXT = RGBColor(0x11, 0x11, 0x11)  # #111111 near-black

# Spanish month abbreviations
_ES_MONTHS = {
    1: "Ene", 2: "Feb", 3: "Mar", 4: "Abr", 5: "May", 6: "Jun",
    7: "Jul", 8: "Ago", 9: "Sep", 10: "Oct", 11: "Nov", 12: "Dic",
}
_EN_MONTHS = {
    1: "Jan", 2: "Feb", 3: "Mar", 4: "Apr", 5: "May", 6: "Jun",
    7: "Jul", 8: "Aug", 9: "Sep", 10: "Oct", 11: "Nov", 12: "Dec",
}

_VALIDATION_COMMENT = (
    "No pudimos verificar este contenido en tu currículum original. "
    "Revísalo antes de enviarlo."
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_run_font(run, size_pt: float, bold: bool = False, color: RGBColor = COLOR_TEXT):
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = color


def _add_paragraph(doc: Document, text: str = "", style: str = "Normal") -> object:
    para = doc.add_paragraph(style=style)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = Pt(12)
    return para


def _configure_page(doc: Document) -> None:
    """Set US Letter page size and margins (DXA units)."""
    section = doc.sections[0]
    section.page_width = Twips(12240)
    section.page_height = Twips(15840)
    section.top_margin = Twips(720)
    section.bottom_margin = Twips(720)
    section.left_margin = Twips(900)
    section.right_margin = Twips(900)


def _add_bullet_paragraph(doc: Document, text: str, flagged: bool = False) -> object:
    """Add a properly formatted bullet paragraph using docx List Bullet style.

    Uses the built-in 'List Bullet' style for proper OOXML numbering definition,
    then overrides indent to match the spec (left=360 DXA, hanging=180 DXA).
    """
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.left_indent = Twips(360)
    para.paragraph_format.first_line_indent = Twips(-180)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    _set_run_font(run, 10)
    if flagged:
        _apply_red_highlight(para)
    return para


def _add_section_header(doc: Document, title: str, flagged: bool = False) -> object:
    """Add a bold ALL-CAPS section header paragraph (no borders — ATS safe)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Twips(120)
    para.paragraph_format.space_after = Twips(60)
    run = para.add_run(title.upper())
    _set_run_font(run, 11, bold=True)
    if flagged:
        _apply_red_highlight(para)
    return para


def _apply_red_highlight(para) -> None:
    """Apply red background highlight to all runs in a paragraph."""
    for run in para.runs:
        rPr = run._r.get_or_add_rPr()
        highlight = OxmlElement("w:highlight")
        highlight.set(qn("w:val"), "red")
        rPr.append(highlight)


def _add_word_comment(doc: Document, para, comment_text: str) -> None:
    """Add a Word comment to a paragraph (Spanish validation warning)."""
    # Word comments require complex XML manipulation; we use a simple annotation
    # via a visible inline note for compatibility across Word versions.
    run = para.add_run(f" [⚠ {comment_text}]")
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.font.size = Pt(8)
    run.font.italic = True


def _render_bold_markdown(para, text: str, size_pt: float) -> None:
    """Render a string with **bold** markdown as alternating bold/normal runs.

    Splits on ** markers — odd-indexed segments (1, 3, …) are bold,
    even-indexed (0, 2, …) are normal weight. Handles the common LLM output
    pattern of **Category**: description without any special casing needed.
    """
    parts = re.split(r'\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        _set_run_font(run, size_pt, bold=(i % 2 == 1))


def _text_is_flagged(text: str, flagged_texts: set[str]) -> bool:
    """Return True if text substring-matches any entry in flagged_texts.

    Checks both directions: the flagged_text may quote only part of the
    content (validator quoting a phrase inside a bullet), or the full
    content may be shorter than the flagged_text. Either way counts as a match.
    """
    t = text.strip().lower()
    return any(ft in t or t in ft for ft in flagged_texts)


# ---------------------------------------------------------------------------
# Filename generation
# ---------------------------------------------------------------------------

def _ascii_safe(text: str) -> str:
    """Strip diacritics and keep only ASCII-safe characters for filenames."""
    import unicodedata
    normalized = unicodedata.normalize("NFD", text)
    return "".join(c for c in normalized if unicodedata.category(c) != "Mn")


def generate_filename(
    candidate_name: str,
    language: str,
    target_company: str,
    date: Optional[datetime] = None,
) -> str:
    """Generate the output filename per the naming convention.

    English: FirstName-LastName-resume-MonYY-CO.docx
    Spanish: FirstName-LastName-CV-MonYY-CO.docx
    """
    if date is None:
        date = datetime.now()

    # Normalize name — strip accents so the filename is ASCII-safe for Supabase Storage
    parts = _ascii_safe(candidate_name.strip()).split()
    if len(parts) >= 2:
        first = parts[0].capitalize()
        last = parts[-1].capitalize()
        name_part = f"{first}-{last}"
    else:
        name_part = _ascii_safe(candidate_name).replace(" ", "-").capitalize()

    # Month abbreviation
    month_map = _ES_MONTHS if language == "es" else _EN_MONTHS
    month_abbr = month_map[date.month]
    year_2d = str(date.year)[-2:]

    # Company code: first 2 ASCII-safe letters, capitalized; fallback 'XX'
    company_safe = _ascii_safe(target_company.strip()) if target_company else ""
    company_code = company_safe[:2].capitalize() or "XX"

    doc_type = "CV" if language == "es" else "resume"
    filename = f"{name_part}-{doc_type}-{month_abbr}{year_2d}-{company_code}.docx"
    return re.sub(r'[^a-zA-Z0-9\-_.]', '', filename)


# ---------------------------------------------------------------------------
# Main document generator
# ---------------------------------------------------------------------------

def generate_docx(
    tailored_resume: TailoredResume,
    flagged_findings: list[dict],
) -> tuple[bytes, str]:
    """Generate a formatted DOCX from a TailoredResume model.

    Args:
        tailored_resume: The tailored resume data.
        flagged_findings: List of dicts with keys 'section' and 'flagged_text',
                          one entry per ValidationFinding that survived repair.
                          For non-experience sections the whole section is
                          highlighted; for experience, individual titles and
                          bullets are matched by text.

    Returns:
        Tuple of (docx_bytes, filename).
    """
    doc = Document()
    _configure_page(doc)

    # Remove default styles spacing
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.font.color.rgb = COLOR_TEXT

    # Sections where the whole block is flagged (summary, skills, education, languages).
    flagged_set = {item["section"] for item in flagged_findings}

    # For experience: normalized flagged texts for per-item matching.
    exp_flagged_texts: set[str] = {
        item["flagged_text"].strip().lower()
        for item in flagged_findings
        if item["section"] == "experience"
    }

    lang = tailored_resume.language

    # ------------------------------------------------------------------
    # Header — Name + Contact
    # ------------------------------------------------------------------
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(0)
    name_run = name_para.add_run(tailored_resume.candidate_name)
    _set_run_font(name_run, 16, bold=True)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_before = Pt(0)
    contact_para.paragraph_format.space_after = Twips(80)
    contact_run = contact_para.add_run(tailored_resume.contact_line)
    _set_run_font(contact_run, 9)

    # ------------------------------------------------------------------
    # Summary
    # ------------------------------------------------------------------
    summary_flagged = "summary" in flagged_set
    summary_label = "RESUMEN" if lang == "es" else "SUMMARY"
    _add_section_header(doc, summary_label, flagged=summary_flagged)

    summary_para = doc.add_paragraph()
    summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary_para.paragraph_format.space_before = Pt(0)
    summary_para.paragraph_format.space_after = Pt(0)
    summary_run = summary_para.add_run(tailored_resume.summary)
    _set_run_font(summary_run, 10)
    if summary_flagged:
        _apply_red_highlight(summary_para)
        _add_word_comment(doc, summary_para, _VALIDATION_COMMENT)

    # ------------------------------------------------------------------
    # Key Skills
    # ------------------------------------------------------------------
    skills_flagged = "skills" in flagged_set
    skills_label = "HABILIDADES CLAVE" if lang == "es" else "KEY SKILLS"
    _add_section_header(doc, skills_label, flagged=skills_flagged)

    skills_para = doc.add_paragraph()
    skills_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    skills_para.paragraph_format.space_before = Pt(0)
    skills_para.paragraph_format.space_after = Pt(0)
    # Render skills with **bold** markdown support (e.g. "**Programming**: Python, SQL")
    skills_text = " | ".join(tailored_resume.skills)
    _render_bold_markdown(skills_para, skills_text, size_pt=10)
    if skills_flagged:
        _apply_red_highlight(skills_para)
        _add_word_comment(doc, skills_para, _VALIDATION_COMMENT)

    # ------------------------------------------------------------------
    # Relevant Work Experience
    # ------------------------------------------------------------------
    exp_label = "EXPERIENCIA PROFESIONAL RELEVANTE" if lang == "es" else "RELEVANT WORK EXPERIENCE"
    # Never highlight the section header — individual items are flagged granularly below.
    _add_section_header(doc, exp_label)

    for job in tailored_resume.experience:
        title_text = f"{job.company} | {job.title}"
        title_flagged = bool(exp_flagged_texts) and _text_is_flagged(title_text, exp_flagged_texts)

        # Job title line with right-aligned date
        job_para = doc.add_paragraph()
        job_para.paragraph_format.space_before = Twips(80)
        job_para.paragraph_format.space_after = Pt(0)
        # Tab stop at right margin (9360 DXA = content width for 0.625" margins)
        tab_stops = job_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Twips(9360), WD_TAB_ALIGNMENT.RIGHT)

        title_run = job_para.add_run(title_text)
        _set_run_font(title_run, 10, bold=True)
        date_run = job_para.add_run(f"\t{job.dates}")
        _set_run_font(date_run, 10, bold=True)

        if title_flagged:
            _apply_red_highlight(job_para)
            _add_word_comment(doc, job_para, _VALIDATION_COMMENT)

        # Bullet points — each checked individually
        for bullet in job.bullets:
            bullet_flagged = bool(exp_flagged_texts) and _text_is_flagged(bullet, exp_flagged_texts)
            bullet_para = _add_bullet_paragraph(doc, bullet, flagged=bullet_flagged)
            if bullet_flagged:
                _add_word_comment(doc, bullet_para, _VALIDATION_COMMENT)

    # ------------------------------------------------------------------
    # Education
    # ------------------------------------------------------------------
    edu_flagged = "education" in flagged_set
    edu_label = "EDUCACIÓN" if lang == "es" else "EDUCATION"
    _add_section_header(doc, edu_label, flagged=edu_flagged)

    for entry in tailored_resume.education:
        edu_para = doc.add_paragraph()
        edu_para.paragraph_format.space_before = Pt(0)
        edu_para.paragraph_format.space_after = Pt(0)
        edu_run = edu_para.add_run(entry)
        _set_run_font(edu_run, 10)
        if edu_flagged:
            _apply_red_highlight(edu_para)

    if edu_flagged and tailored_resume.education:
        _add_word_comment(
            doc, doc.paragraphs[-1], _VALIDATION_COMMENT
        )

    # ------------------------------------------------------------------
    # Languages
    # ------------------------------------------------------------------
    lang_flagged = "languages" in flagged_set
    lang_label = "IDIOMAS" if lang == "es" else "LANGUAGES"
    _add_section_header(doc, lang_label, flagged=lang_flagged)

    lang_para = doc.add_paragraph()
    lang_para.paragraph_format.space_before = Pt(0)
    lang_para.paragraph_format.space_after = Pt(0)
    # First language entry is bold; remaining entries are regular (per formatting.md)
    languages_line = tailored_resume.languages_line
    if " | " in languages_line:
        sep_idx = languages_line.index(" | ")
        first_run = lang_para.add_run(languages_line[:sep_idx])
        _set_run_font(first_run, 10, bold=True)
        rest_run = lang_para.add_run(languages_line[sep_idx:])
        _set_run_font(rest_run, 10)
    else:
        lang_run = lang_para.add_run(languages_line)
        _set_run_font(lang_run, 10, bold=True)
    if lang_flagged:
        _apply_red_highlight(lang_para)
        _add_word_comment(doc, lang_para, _VALIDATION_COMMENT)

    # ------------------------------------------------------------------
    # Serialize to bytes
    # ------------------------------------------------------------------
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    filename = generate_filename(
        tailored_resume.candidate_name,
        tailored_resume.language,
        tailored_resume.target_company,
    )

    return docx_bytes, filename
